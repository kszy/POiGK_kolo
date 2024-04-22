from utils import random_strings
import json
from docx import Document
from docx.shared import Inches, Pt
import numpy as np
import random


class QuizGeneration:
    def __init__(
        self,
        num_students,
        num_questions,
        questions_path="questions.json",
        out_quiz_path="quiz.docx",
        out_answers_path="answers.docx",
    ):
        self.quiz_keys = random_strings(2, num_students)
        self.num_questions = num_questions
        self.question_set = self.load_questions(questions_path)
        print(len(self.question_set))
        self.quiz_document = Document()
        self.out_quiz_path = out_quiz_path
        self.answer_document = Document()
        self.document_settings()
        self.out_answers_path = out_answers_path
        self.generate_quizzes()

    def document_settings(self):
        style = self.quiz_document.styles["Normal"]
        font = style.font
        font.size = Pt(9)
        self.quiz_document.sections[0].left_margin = Inches(0.4)
        self.quiz_document.sections[0].right_margin = Inches(0.4)
        self.quiz_document.sections[0].top_margin = Inches(0.4)
        self.quiz_document.sections[0].bottom_margin = Inches(0.4)

    def load_questions(self, questions_path):
        return json.load(open(questions_path))

    def generate_quizzes(self):
        for key in self.quiz_keys:
            self.generate_quiz(key)
        self.quiz_document.save(self.out_quiz_path)
        self.answer_document.save(self.out_answers_path)

    def generate_quiz(self, key):
        h = self.quiz_document.add_heading(
            "Kolokwium z Programowania Obiektowego i Grafiki Komputerowej", 2
        )
        h.add_run(f" ({key})")
        self.answer_document.add_paragraph(key)
        answer_table = self.gen_answer_table(self.answer_document)
        question_indices = np.random.choice(
            len(self.question_set), self.num_questions, replace=False
        )
        print(question_indices)
        for q_num, q_idx in enumerate(question_indices):
            self.generate_question(q_num, q_idx, answer_table)
        self.gen_answer_table(self.quiz_document)
        self.quiz_document.add_page_break()

    def generate_question(self, q_num, q_idx, answer_table):
        p = self.quiz_document.add_paragraph()
        p.add_run(f'{q_num+1} {self.question_set[q_idx]["question"]}').bold = True
        answer_order = np.random.permutation(len(self.question_set[q_idx]["answers"]))
        answer_table.cell(1, q_num).text = chr(np.where(answer_order == 0)[0][0] + 65)
        for i, answer in enumerate(
            [self.question_set[q_idx]["answers"][i] for i in answer_order]
        ):
            p.add_run(f"\n{chr(i + 65)})").bold = True
            p.add_run(f" {answer}")

    def gen_answer_table(self, doc):
        table = doc.add_table(rows=2, cols=self.num_questions)
        table.style = "Table Grid"
        for i in range(self.num_questions):
            table.cell(0, i).text = str(i + 1)
        return table


if __name__ == "__main__":
    QuizGeneration(120, 17)
